from flask import Flask, request, Blueprint, jsonify, send_file
from langchain_ollama import OllamaLLM
from langchain_core.prompts import ChatPromptTemplate
import docx
from io import BytesIO
from flask_cors import CORS

chatbot_letter = Blueprint('chatbot_letter', __name__)

# Fonction pour extraire le texte d'un fichier DOCX
# Fonction pour extraire le texte d'un fichier DOCX
def extract_text_from_docx(docx_file):
    try:
        doc = docx.Document(docx_file)
        text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
        return text
    except Exception as e:
        print("Erreur lors de l'extraction du fichier DOCX :", str(e))
        return None
 
# Analyse et suggestions d'amélioration
analyze_template = """Tu es un expert en rédaction de lettres de motivation. Voici une lettre de motivation et la description d'une offre d'emploi. Identifie les points faibles de la lettre et propose des suggestions d'amélioration détaillées et réalistes adaptées à l'offre d'emploi.
 
Lettre de motivation : {context}
Description de l'offre : {job_description}
 
Réponse :"""
analyze_model = OllamaLLM(model="llama3.2")
analyze_prompt = ChatPromptTemplate.from_template(analyze_template)
analyze_chain = analyze_prompt | analyze_model
 
# Génération d'une lettre optimisée
generate_template = """Tu es un expert en rédaction de lettres de motivation. Basé sur la lettre ci-dessous et la description d'une offre d'emploi, rédige une nouvelle lettre optimisée. La réponse doit être claire, structurée et adaptée au poste.
 
Lettre de motivation : {context}
Description de l'offre : {job_description}
 
Nouvelle lettre :"""
generate_model = OllamaLLM(model="llama3.2")
generate_prompt = ChatPromptTemplate.from_template(generate_template)
generate_chain = generate_prompt | generate_model
 
# Fonction pour générer une lettre optimisée
def generate_optimized_letter(letter_text, job_description):
    result = generate_chain.invoke({"context": letter_text, "job_description": job_description})
    return result.strip()
 
# Fonction pour analyser une lettre
def analyze_letter(letter_text, job_description):
    result = analyze_chain.invoke({"context": letter_text, "job_description": job_description})
    return result.strip()
 
# Fonction pour sauvegarder un texte dans un fichier Word
def save_text_to_word(text, filename="lettre_de_motivation.docx"):
    doc = docx.Document()
    doc.add_paragraph(text)
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
 
@chatbot_letter.route('/analyze-letter', methods=['POST'])
def analyze():
    try:
        # Logs pour déboguer les données reçues
        print("Debug: Requête reçue :", request)
        print("Debug: Fichiers reçus :", request.files)
        print("Debug: Données reçues :", request.form)
 
        # Récupérer le fichier et la description
        file = request.files.get('file')
        job_description = request.form.get('job_description')
 
        if not file or not job_description:
            return jsonify({"error": "Le fichier et la description de l'offre sont requis."}), 400
 
        # Extraire le contenu de la lettre
        letter_text = extract_text_from_docx(file)
        if not letter_text:
            return jsonify({"error": "Impossible de lire le fichier DOCX."}), 400
 
        # Analyse de la lettre
        feedback = analyze_letter(letter_text, job_description)
        if not feedback.strip():
            return jsonify({"error": "Aucune suggestion trouvée. Vérifiez les entrées."}), 500
 
        return jsonify({"feedback": feedback})
    except Exception as e:
        print("Erreur serveur :", str(e))
        return jsonify({"error": "Erreur inattendue : " + str(e)}), 500
 
 
@chatbot_letter.route('/generate-letter', methods=['POST'])
def generate():
    try:
        # Logs pour déboguer les données reçues
        print("Debug: Requête reçue :", request)
        print("Debug: Fichiers reçus :", request.files)
        print("Debug: Données reçues :", request.form)
 
        file = request.files.get('file')
        job_description = request.form.get('job_description')
 
        if not file or not job_description:
            return jsonify({"error": "Le fichier et la description de l'offre sont requis."}), 400
 
        letter_text = extract_text_from_docx(file)
        if not letter_text:
            return jsonify({"error": "Impossible de lire le fichier DOCX."}), 400
 
        new_letter = generate_optimized_letter(letter_text, job_description)
        if not new_letter.strip():
            return jsonify({"error": "Aucune lettre générée. Vérifiez les entrées."}), 500
 
        return jsonify({"new_letter": new_letter})
    except Exception as e:
        print("Erreur serveur :", str(e))
        return jsonify({"error": "Erreur inattendue : " + str(e)}), 500
 
@chatbot_letter.route('/download-letter', methods=['POST'])
def download_letter():
    try:
        data = request.json
        new_letter = data.get("new_letter")
 
        if not new_letter:
            return jsonify({"error": "La lettre optimisée est requise pour le téléchargement."}), 400
 
        # Sauvegarder la lettre dans un fichier Word
        word_file = save_text_to_word(new_letter)
        return send_file(
            word_file,
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            as_attachment=True,
            download_name="lettre_optimisee.docx"
        )
    except Exception as e:
        print("Erreur serveur :", str(e))
        return jsonify({"error": "Erreur inattendue : " + str(e)}), 500